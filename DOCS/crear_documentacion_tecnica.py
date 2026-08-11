from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_BREAK
from datetime import date
from pathlib import Path

OUT = Path(__file__).with_name("Documentacion_Tecnica_Retorno360Tacna.docx")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "5B6573"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"


def set_font(run, name="Calibri", size=11, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for cell, text in zip(hdr.cells, headers):
        shade(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_font(r, size=9.5, color=INK, bold=True)
    for values in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, values):
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            set_font(r, size=9.2, color="222222")
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="Normal")
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_font(r, bold=True)
        r = p.add_run(text[len(bold_prefix):])
        set_font(r)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_font(r)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    r = p.add_run(text)
    set_font(r, size={1: 16, 2: 13, 3: 12}[level], color={1: BLUE, 2: BLUE, 3: DARK_BLUE}[level], bold=True)
    return p


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade(cell, "F4F6F9")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title + " ")
    set_font(r, size=10.5, color=INK, bold=True)
    r = p.add_run(text)
    set_font(r, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        s = styles[name]
        s.font.name = "Calibri"
        s._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        s._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        s.font.size = Pt(size)
        s.font.color.rgb = RGBColor.from_string(color)
        s.font.bold = True
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("Retorno360 Tacna | Documentación técnica")
    set_font(r, size=8.5, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("Documento interno - generado el 7 de agosto de 2026")
    set_font(r, size=8.5, color=MUTED)


def build():
    doc = Document()
    configure_document(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(30)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("DOCUMENTACIÓN TÉCNICA")
    set_font(r, size=13, color=BLUE, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Retorno360 Tacna")
    set_font(r, size=28, color=INK, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("Aplicación de escritorio para gestión de retorno aduanero, conciliación fiscal e inventarios")
    set_font(r, size=13, color=MUTED)
    add_table(doc, ["Elemento", "Detalle"], [
        ("Tipo de aplicación", "Aplicación de escritorio Windows (WinForms)."),
        ("Framework objetivo", ".NET 10, destino net10.0-windows."),
        ("Lenguaje", "C# con nulabilidad habilitada."),
        ("Fecha del documento", "7 de agosto de 2026."),
        ("Alcance", "Código fuente de la solución Retorno360Tacna y sus dependencias declaradas."),
    ], [2700, 6660])
    add_callout(doc, "Propósito.", "Este documento describe cómo está organizada la aplicación, qué funciones atiende cada capa y cuáles son sus integraciones principales. No incluye valores de credenciales ni secretos operativos.")

    add_heading(doc, "1. Visión general", 1)
    add_body(doc, "Retorno360 Tacna centraliza el cálculo y análisis de retorno aduanero para operaciones de comercio exterior. La aplicación consulta fuentes SQL Server, consolida información de pedimentos, calcula indicadores de retorno, concilia IGI e IVA, genera reportes PDF y Excel, y permite publicar información seleccionada hacia un portal web e inventarios.")
    add_body(doc, "La solución sigue una organización por responsabilidades: formularios de WinForms para interacción, servicios para reglas de negocio e integración, modelos para transportar datos y una capa de conexión para SQL Server.")

    add_heading(doc, "2. Arquitectura", 1)
    add_body(doc, "La arquitectura es una aplicación cliente de escritorio con servicios internos. Aunque no adopta una separación estricta de dominio e infraestructura, el código se agrupa de forma clara por carpetas y responsabilidades.")
    add_table(doc, ["Capa", "Responsabilidad", "Ejemplos"], [
        ("Presentación", "Pantallas, eventos, validación de entrada, renderizado de tablas, gráficas y navegación.", "FORMS: Login, MainMenu, FrmRetorno, FrmReportes, FrmReportesInventario."),
        ("Aplicación / servicios", "Orquesta consultas, cálculos, generación de reportes e integraciones externas.", "SERVICES: RetornoService, ReporteIGIService, PdfGeneradorService, PortalWebService."),
        ("Dominio / modelos", "Representa usuarios, conexiones, resultados, pedimentos, catálogo e inventario.", "MODELS: Usuario, ReporteIGI, ConexionInfo, SesionCalculoInventario."),
        ("Infraestructura", "Conexiones SQL Server, PostgreSQL, Cloudflare R2, archivos, secretos y notificaciones.", "CNX/Conexion, CloudflareR2Service, SecretStoreService, ConfiguracionService."),
    ], [1500, 3900, 3960])
    add_heading(doc, "2.1 Flujo principal de datos", 2)
    add_bullet(doc, "El usuario inicia sesión y selecciona una razón social, base de datos y periodo de trabajo.")
    add_bullet(doc, "Los formularios invocan servicios para consultar SQL Server y consolidar datos de pedimentos e inventarios.")
    add_bullet(doc, "Los servicios devuelven modelos y DataTables; la interfaz los muestra en tablas y gráficas.")
    add_bullet(doc, "Los resultados pueden exportarse a PDF/Excel, subirse al portal PostgreSQL o complementarse con archivos almacenados en R2.")

    add_heading(doc, "3. Framework, bibliotecas y plataforma", 1)
    add_table(doc, ["Tecnología", "Uso dentro de la aplicación"], [
        (".NET 10 / WinForms", "Base de la aplicación de escritorio, formularios, controles, eventos y ciclo de vida."),
        ("Microsoft.Data.SqlClient", "Acceso a SQL Server para configuración, usuarios, catálogos y fuentes operativas."),
        ("Npgsql", "Conexión a PostgreSQL/Railway para el portal web y persistencia de resultados publicados."),
        ("LiveChartsCore + SkiaSharp", "Gráficas de retorno, IGI, IVA, inventarios y visualización de datos."),
        ("QuestPDF", "Generación de reportes PDF con tablas, resúmenes y gráficas."),
        ("ClosedXML / OpenXML", "Lectura, generación y exportación de archivos Excel."),
        ("AWS SDK S3", "Cliente compatible con Cloudflare R2 para listar, descargar y eliminar archivos."),
    ], [3000, 6360])

    add_heading(doc, "4. Módulos funcionales", 1)
    add_table(doc, ["Módulo", "Funciones principales", "Componentes relacionados"], [
        ("Autenticación y sesión", "Validar usuarios, resolver roles, cargar conexiones activas y abrir el menú principal.", "Login, LoginService, Usuario, Rol."),
        ("Retorno aduanero", "Calcular porcentaje de retorno, validar pedimentos, consultar importaciones/exportaciones y generar resultados.", "FrmRetorno, RetornoService, CalculadoraRetorno."),
        ("Reportes IGI/IVA", "Conciliar valores pagados y calculados, separar formas de pago, mostrar detalle y exportar reportes.", "FrmReportes, ReporteIGIService, ReporteServiceBase, ReporteIGI."),
        ("Inventarios", "Calcular sesiones de inventario, revisar indicadores, consultar archivos por mes y producir reportes.", "FrmCalculoInventarios, FrmReportesInventario, SesionCalculoInventario."),
        ("Catálogo de partes", "Consultar materia prima/BOM, filtrar por fecha y representar resultados en tablas y gráficas.", "FrmCatalogoPartes, CatalogoPartesService, ParteBOMCompleto."),
        ("Configuración", "Administrar usuarios, conexiones, parámetros y respaldo de configuración.", "FrmConfiguracion, ConfiguracionService, ConexionInfo."),
        ("Documentos y anexos", "Navegar archivos, descargar anexos, revisar calendarios y apoyar solicitudes operativas.", "FrmAnexos, FrmSolicitud, CloudflareR2Service."),
        ("Portal web", "Publicar resultados de cálculo, reportes y datos de inventario hacia PostgreSQL.", "PortalWebService, ConfiguracionService, SecretStoreService."),
    ], [1700, 4700, 2960])

    add_heading(doc, "5. Capas y componentes clave", 1)
    add_heading(doc, "5.1 Presentación: FORMS", 2)
    add_body(doc, "Los formularios WinForms coordinan la interacción con el usuario. Contienen controles, eventos y lógica de presentación para filtros, validaciones, DataGridView, barras de progreso y gráficas. MainMenu funciona como contenedor de navegación; Login controla el acceso inicial; los formularios especializados atienden retorno, reportes, inventarios, configuración y anexos.")
    add_heading(doc, "5.2 Servicios: SERVICES", 2)
    add_table(doc, ["Servicio", "Responsabilidad"], [
        ("RetornoService", "Obtiene información operativa, resuelve conexiones externas y calcula datos de retorno."),
        ("ReporteIGIService / Extensión", "Construye conciliaciones IGI/IVA, agrupa periodos y prepara detalles de reporte."),
        ("ReporteServiceBase", "Comparte operaciones de consulta y resolución de conexiones entre servicios de reportes."),
        ("PdfGeneradorService", "Compone documentos PDF y elementos visuales de los reportes."),
        ("PortalWebService", "Inserta resultados y detalles en PostgreSQL mediante operaciones asíncronas y transacciones."),
        ("CloudflareR2Service", "Lista carpetas/archivos y descarga o elimina objetos en almacenamiento compatible con S3."),
        ("ConfiguracionService", "Construye y valida la conexión del portal a partir de variables de entorno o secretos locales."),
        ("SecretStoreService", "Guarda secretos locales protegidos con DPAPI para el usuario actual de Windows."),
    ], [3300, 6060])
    add_heading(doc, "5.3 Modelos: MODELS", 2)
    add_body(doc, "Los modelos representan el contrato de datos entre servicios y formularios. Incluyen identidad y permisos (Usuario, Rol), configuración de conectividad (ConexionInfo, ConexionExternaInfo), negocio aduanero (ReporteIGI, DatoDetalleIGI, PedimentosPorRazon), materia prima (MateriaPrimaBOM, ParteBOMCompleto, ComponenteBOM) e inventario (SesionCalculoInventario, UcMesInventario).")
    add_heading(doc, "5.4 Conexiones e infraestructura", 2)
    add_body(doc, "La clase Conexion encapsula la construcción y apertura de conexiones SQL Server. Los servicios de portal usan Npgsql para PostgreSQL y el servicio R2 usa el SDK de AWS para Cloudflare R2. Esta distribución permite a la aplicación leer fuentes operativas y publicar resultados hacia sistemas complementarios.")

    add_heading(doc, "6. Persistencia e integraciones", 1)
    add_table(doc, ["Sistema", "Finalidad", "Modo de acceso"], [
        ("SQL Server", "Datos maestros, usuarios, configuraciones, pedimentos, glosa y fuentes operativas.", "Microsoft.Data.SqlClient; conexiones principal y externas."),
        ("PostgreSQL / Railway", "Portal web: resultados de retorno, reportes e inventarios publicados.", "Npgsql; cadena tomada de variable de entorno o almacén DPAPI."),
        ("Cloudflare R2", "Repositorio de archivos, calendarios, anexos e información asociada a clientes.", "AWS SDK S3 compatible con API S3."),
        ("GitHub Releases", "Actualización de la aplicación instalada mediante componente actualizador independiente.", "Cliente HTTP/Octokit en el actualizador distribuido."),
    ], [2300, 4400, 2660])

    add_heading(doc, "7. Procesos representativos", 1)
    add_heading(doc, "7.1 Cálculo de retorno", 2)
    add_bullet(doc, "El usuario define razón social, base de datos, periodo y opciones de cálculo.")
    add_bullet(doc, "FrmRetorno solicita la información a RetornoService y recibe importaciones, exportaciones y pedimentos validados.")
    add_bullet(doc, "CalculadoraRetorno determina el porcentaje y los resultados se presentan en controles y gráficas.")
    add_bullet(doc, "El usuario puede exportar el resultado a PDF y, cuando procede, publicarlo al portal web.")
    add_heading(doc, "7.2 Reporte y conciliación IGI/IVA", 2)
    add_bullet(doc, "FrmReportes solicita a ReporteIGIService el detalle de pedimentos y formas de pago para un intervalo de fechas.")
    add_bullet(doc, "El servicio cruza fuentes de cliente y glosa, calcula diferencias y agrupa resultados por mes y forma de pago.")
    add_bullet(doc, "La interfaz genera tablas, gráficas y una salida PDF con resumen ejecutivo y detalle operativo.")
    add_heading(doc, "7.3 Inventarios y documentos", 2)
    add_bullet(doc, "Los módulos de inventarios conservan una sesión de cálculo y permiten revisar resultados por mes o razón social.")
    add_bullet(doc, "Los módulos de anexos/reportes de inventario consultan R2 para navegar y descargar archivos asociados.")

    add_heading(doc, "8. Seguridad y operación", 1)
    add_callout(doc, "Observación de seguridad.", "La aplicación cuenta con un almacén DPAPI para secretos locales, pero el código revisado contiene configuraciones sensibles incrustadas. Estos valores no se documentan aquí; deben migrarse a un almacén de secretos y rotarse antes de cualquier distribución.")
    add_table(doc, ["Área", "Estado observado", "Recomendación"], [
        ("Secretos", "Existe SecretStoreService con DPAPI, pero hay configuraciones sensibles fuera de este mecanismo.", "Rotar valores expuestos y usar únicamente secretos por usuario/entorno."),
        ("Contraseñas", "La autenticación actual usa hash SHA-256 simple.", "Migrar a Argon2id, bcrypt o PBKDF2 con sal única y costo configurado."),
        ("Conexiones", "Las consultas revisadas usan parámetros en los puntos principales; algunas conexiones confían en el certificado del servidor.", "Validar certificados TLS en producción y sustituir AddWithValue por parámetros tipados donde sea relevante."),
        ("Actualización", "El ejecutable distribuido y el flujo de actualización requieren endurecimiento.", "Firmar binarios y verificar la integridad/autenticidad de cada paquete de actualización."),
    ], [1900, 3900, 3560])

    add_heading(doc, "9. Calidad técnica y mantenimiento", 1)
    add_body(doc, "La solución compiló sin errores durante la revisión, pero se registraron 157 advertencias. Predominan advertencias de nulabilidad, APIs obsoletas de visualización, miembros que ocultan métodos heredados y compatibilidad de paquetes. No se detectaron proyectos de pruebas automatizadas en la solución.")
    add_bullet(doc, "Priorizar correcciones de nulabilidad en formularios de reportes, inventarios y catálogo de partes para prevenir errores en tiempo de ejecución.")
    add_bullet(doc, "Actualizar dependencias con vulnerabilidades o restricciones incompatibles y fijar versiones estables donde hoy se usan comodines.")
    add_bullet(doc, "Separar gradualmente la lógica de negocio y acceso a datos que aún reside en formularios extensos.")
    add_bullet(doc, "Crear pruebas unitarias para cálculos de retorno, conciliaciones IGI/IVA, transformaciones de datos y generación de resúmenes.")

    add_heading(doc, "10. Estructura del proyecto", 1)
    add_table(doc, ["Ruta", "Contenido"], [
        ("FORMS", "Formularios WinForms, diseñadores y recursos visuales."),
        ("SERVICES", "Lógica de aplicación, reportes, almacenamiento, portal y configuración."),
        ("MODELS", "Entidades, DTOs y clases auxiliares del dominio."),
        ("CNX", "Conexión y utilidades para SQL Server."),
        ("HELPERS", "Notificaciones, registro de errores, copia de DataGridView y utilidades de interfaz."),
        ("Properties / Resources", "Recursos embebidos, configuración y metadatos de la aplicación."),
        ("DOCS", "Documentación técnica y material de apoyo del repositorio."),
    ], [2700, 6660])

    add_heading(doc, "11. Recomendaciones de evolución", 1)
    for item in [
        "Eliminar y rotar secretos existentes; establecer una estrategia única de configuración por entorno.",
        "Fortalecer la autenticación y actualizar los hashes de contraseñas existentes de forma progresiva.",
        "Introducir una capa de repositorios o adaptadores para desacoplar SQL Server, PostgreSQL y R2 de los formularios.",
        "Reducir el tamaño de los formularios de reportes mediante controladores/presentadores y servicios especializados.",
        "Incorporar pruebas unitarias y pruebas de integración con bases de datos de prueba.",
        "Implantar CI que compile, ejecute pruebas, audite paquetes y firme los artefactos de distribución.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "12. Conclusión", 1)
    add_body(doc, "Retorno360 Tacna es una aplicación de escritorio funcionalmente amplia, orientada a procesos aduaneros, conciliación fiscal e inventarios. Su base tecnológica moderna (.NET 10) y sus integraciones cubren las necesidades principales del negocio. La siguiente etapa de madurez debe centrarse en seguridad de secretos y credenciales, confiabilidad de dependencias, reducción de advertencias y automatización de pruebas y despliegues.")

    doc.core_properties.title = "Documentación Técnica - Retorno360 Tacna"
    doc.core_properties.subject = "Arquitectura, capas, funciones y tecnologías"
    doc.core_properties.author = "Retorno360 Tacna"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
